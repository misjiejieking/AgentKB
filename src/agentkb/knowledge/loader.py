""".md / .txt / .pdf / .docx / .csv / .json 文件加载与校验——统一注册机制。"""

from __future__ import annotations

import csv
import json
from abc import ABC, abstractmethod
from io import StringIO
from pathlib import Path

from langchain_community.document_loaders import TextLoader
from langchain_core.documents import Document
from loguru import logger

from agentkb.utils.exceptions import KnowledgeBaseError
from agentkb.config.settings import Settings
from agentkb.multimodal.vision import VisionService


class BaseFileLoader(ABC):
    """文件加载器基类——每个子类负责一种文件格式。"""

    @property
    @abstractmethod
    def extensions(self) -> set[str]:
        """支持的文件扩展名集合，含点号（如 {'.pdf', '.docx'}）。"""
        ...

    @abstractmethod
    def load(self, file_path: Path) -> list[Document]:
        """加载文件为 LangChain Document 列表。"""
        ...


# ══════════════════════════════════════════════════════════════
#  内置 Loader 实现
# ══════════════════════════════════════════════════════════════

class MarkdownLoader(BaseFileLoader):
    """.md / .txt 文本加载器。"""

    @property
    def extensions(self) -> set[str]:
        return {".md", ".txt"}

    def load(self, file_path: Path) -> list[Document]:
        loader = TextLoader(str(file_path), encoding="utf-8")
        docs = loader.load()
        for doc in docs:
            doc.metadata["source"] = file_path.name
        return docs


class PdfLoader(BaseFileLoader):
    """.pdf 加载器，保留页面文本并可分析扫描页和图表。"""

    def __init__(
        self,
        settings: Settings,
        vision_service: VisionService | None = None,
    ) -> None:
        self.settings = settings
        self.vision_service = vision_service

    @property
    def extensions(self) -> set[str]:
        return {".pdf"}

    def load(self, file_path: Path) -> list[Document]:
        import fitz  # type: ignore[import-untyped]  # PyMuPDF 未提供类型声明
        docs = []
        try:
            with fitz.open(str(file_path)) as doc:
                visual_pages = 0
                for page_num, page in enumerate(doc):
                    text = page.get_text().strip()
                    visual_text = ""
                    needs_visual = (
                        self.settings.vision_pdf_visual_analysis
                        and visual_pages < self.settings.vision_pdf_max_pages
                        and (len(text) < 80 or bool(page.get_images(full=True)))
                    )
                    if needs_visual:
                        pixmap = page.get_pixmap(
                            matrix=fitz.Matrix(1.5, 1.5),
                            alpha=False,
                        )
                        image = pixmap.tobytes("jpeg")
                        service = self.vision_service or VisionService(
                            settings=self.settings
                        )
                        analysis = service.analyze(
                            image,
                            prompt=(
                                f"这是 PDF《{file_path.name}》第 {page_num + 1} 页。"
                                "请提取页面文字，并解释图表、流程图、公式和版面关系；"
                                "输出可独立检索的中文描述。"
                            ),
                        )
                        visual_text = analysis.description
                        visual_pages += 1

                    page_content = "\n\n".join(
                        part for part in (text, visual_text) if part
                    )
                    if page_content:
                        docs.append(Document(
                            page_content=page_content,
                            metadata={
                                "source": file_path.name,
                                "page_number": page_num + 1,
                                "total_pages": len(doc),
                                "visual_analysis": bool(visual_text),
                            },
                        ))
        except Exception as e:
            raise KnowledgeBaseError(f"PDF 解析失败 '{file_path.name}': {e}") from e

        if not docs:
            raise KnowledgeBaseError(f"PDF 文件中无有效文本内容: {file_path.name}")
        return docs


class ImageLoader(BaseFileLoader):
    """图片知识加载器，将视觉内容转换为可检索文本。"""

    def __init__(self, vision_service: VisionService) -> None:
        self.vision_service = vision_service

    @property
    def extensions(self) -> set[str]:
        return {".jpg", ".jpeg", ".png", ".webp"}

    def load(self, file_path: Path) -> list[Document]:
        analysis = self.vision_service.analyze(file_path.read_bytes())
        return [Document(
            page_content=analysis.description,
            metadata={
                "source": file_path.name,
                "media_type": analysis.media_type,
                "visual_analysis": True,
            },
        )]


class DocxLoader(BaseFileLoader):
    """.docx 加载器（python-docx），按段落提取，保留标题样式信息。"""

    @property
    def extensions(self) -> set[str]:
        return {".docx"}

    def load(self, file_path: Path) -> list[Document]:
        from docx import Document as DocxDocument
        from docx.enum.style import WD_STYLE_TYPE

        try:
            doc = DocxDocument(str(file_path))
        except Exception as e:
            raise KnowledgeBaseError(f"DOCX 解析失败 '{file_path.name}': {e}") from e

        paragraphs: list[tuple[str, bool]] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            is_heading = para.style and para.style.type == WD_STYLE_TYPE.PARAGRAPH and "heading" in (para.style.name or "").lower()
            paragraphs.append((text, bool(is_heading)))

        if not paragraphs:
            raise KnowledgeBaseError(f"DOCX 文件中无有效文本内容: {file_path.name}")

        # 整个文档作为一个 Document，段落结构保留到 metadata
        full_text = "\n\n".join(
            (f"## {text}" if is_heading else text)
            for text, is_heading in paragraphs
        )
        return [Document(
            page_content=full_text,
            metadata={"source": file_path.name, "paragraph_count": len(paragraphs)},
        )]


class CsvLoader(BaseFileLoader):
    """.csv 加载器，按行处理，表头作为 metadata。"""

    @property
    def extensions(self) -> set[str]:
        return {".csv"}

    def load(self, file_path: Path) -> list[Document]:
        try:
            content = file_path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            content = file_path.read_text(encoding="gbk")

        reader = csv.reader(StringIO(content))
        headers = next(reader, None)
        if not headers:
            raise KnowledgeBaseError(f"CSV 文件无表头: {file_path.name}")

        # 每 50 行合并为一个 Document
        docs = []
        batch_lines = []
        row_count = 0
        for row in reader:
            batch_lines.append(", ".join(row))
            row_count += 1
            if len(batch_lines) >= 50:
                docs.append(Document(
                    page_content="\n".join(batch_lines),
                    metadata={"source": file_path.name, "headers": headers},
                ))
                batch_lines = []

        if batch_lines:
            docs.append(Document(
                page_content="\n".join(batch_lines),
                metadata={"source": file_path.name, "headers": headers},
            ))

        if not docs:
            raise KnowledgeBaseError(f"CSV 文件中无数据行: {file_path.name}")
        return docs


class JsonLoader(BaseFileLoader):
    """.json 加载器，数组按元素展开，对象整体作为文本。"""

    @property
    def extensions(self) -> set[str]:
        return {".json"}

    def load(self, file_path: Path) -> list[Document]:
        try:
            data = json.loads(file_path.read_text(encoding="utf-8"))
        except json.JSONDecodeError as e:
            raise KnowledgeBaseError(f"JSON 解析失败 '{file_path.name}': {e}") from e

        if isinstance(data, list):
            docs = []
            for i, item in enumerate(data):
                if isinstance(item, dict):
                    text = json.dumps(item, ensure_ascii=False, indent=2)
                else:
                    text = str(item)
                docs.append(Document(
                    page_content=text,
                    metadata={"source": file_path.name, "index": i},
                ))
            if not docs:
                raise KnowledgeBaseError(f"JSON 数组为空: {file_path.name}")
            return docs

        # 单个对象：格式化输出
        text = json.dumps(data, ensure_ascii=False, indent=2) if isinstance(data, dict) else str(data)
        return [Document(page_content=text, metadata={"source": file_path.name})]


# ══════════════════════════════════════════════════════════════
#  FileLoader — 统一入口
# ══════════════════════════════════════════════════════════════

class FileLoader:
    """统一文件加载入口，按扩展名分发到对应 Loader。"""

    def __init__(
        self,
        settings: Settings | None = None,
        vision_service: VisionService | None = None,
    ) -> None:
        settings = settings or Settings.load()
        vision_service = vision_service or VisionService(settings=settings)
        self._loaders: dict[str, BaseFileLoader] = {}
        for loader in (
            MarkdownLoader(),
            PdfLoader(settings, vision_service),
            DocxLoader(),
            CsvLoader(),
            JsonLoader(),
            ImageLoader(vision_service),
        ):
            for ext in loader.extensions:
                self._loaders[ext] = loader

    # ── 加载 ──────────────────────────────────────────────────

    def load(self, file_path: str | Path) -> list[Document]:
        path = Path(file_path)
        ext = path.suffix.lower()
        loader = self._loaders.get(ext)
        if not loader:
            raise KnowledgeBaseError(f"不支持的文件类型: {ext}")
        docs = loader.load(path)
        logger.debug(f"已加载 {len(docs)} 个文档，来源: {path.name}（{ext}）")
        return docs
